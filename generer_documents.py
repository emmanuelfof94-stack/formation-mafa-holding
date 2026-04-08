#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de génération des documents Gestion Qualité
- QR Codes des agents
- Règlement Intérieur (Word)
- Fichier de pointage (Excel)
- Registre des agents (Excel)
"""

import json
import os
import qrcode
from datetime import datetime, date
from pathlib import Path

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
QR_DIR = BASE_DIR / "qrcodes"
EXPORT_DIR = BASE_DIR / "exports"

# ─────────────────────────────────────────────
# 1. GÉNÉRATION DES QR CODES
# ─────────────────────────────────────────────
def generer_qrcodes():
    with open(DATA_DIR / "agents.json", encoding="utf-8") as f:
        agents = json.load(f)

    for agent in agents:
        data = f"BADGE|{agent['id']}|{agent['nom']}|{agent['prenom']}"
        qr = qrcode.QRCode(version=2, box_size=10, border=4,
                           error_correction=qrcode.constants.ERROR_CORRECT_H)
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        path = QR_DIR / f"QR_{agent['id']}_{agent['nom']}.png"
        img.save(str(path))
        print(f"  ✓ QR Code généré : {path.name}")

# ─────────────────────────────────────────────
# 2. RÈGLEMENT INTÉRIEUR (WORD)
# ─────────────────────────────────────────────
def generer_reglement_word():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Marges
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── En-tête ──
    header = doc.add_heading("", 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("RÈGLEMENT INTÉRIEUR")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sous_titre.add_run("Service Gestion Qualité")
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(f"Date d'entrée en vigueur : {date.today().strftime('%d/%m/%Y')}").font.size = Pt(10)

    doc.add_paragraph()

    # ── Introduction ──
    intro = doc.add_paragraph()
    intro.add_run(
        "Le présent règlement intérieur s'applique à l'ensemble du personnel affecté au service. "
        "Son respect est obligatoire et conditionne le bon fonctionnement collectif de notre environnement de travail. "
        "Tout manquement pourra faire l'objet d'une mesure disciplinaire."
    ).font.size = Pt(11)

    doc.add_paragraph()

    # ── Règles ──
    regles = [
        {
            "num": "Art. 1",
            "titre": "Utilisation du téléphone personnel",
            "contenu": [
                "L'utilisation du téléphone personnel est strictement interdite durant les heures de travail.",
                "Les appareils doivent être mis en mode silencieux et rangés dès la prise de poste.",
                "En cas d'urgence personnelle, l'agent doit en informer son supérieur et s'isoler dans un espace dédié.",
                "Tout manquement répété pourra entraîner une sanction disciplinaire."
            ]
        },
        {
            "num": "Art. 2",
            "titre": "Système de pointage — Badge QR Code",
            "contenu": [
                "Chaque agent est tenu de badger à l'arrivée (matin) et au départ (soir) via le QR code personnel remis lors de son intégration.",
                "Le badgeage s'effectue sur la borne installée à l'entrée du service.",
                "Tout oubli de badgeage doit être signalé immédiatement au responsable.",
                "Les données de pointage servent de base au suivi des présences et à l'établissement des rapports mensuels."
            ]
        },
        {
            "num": "Art. 3",
            "titre": "Code vestimentaire",
            "contenu": [
                "Une tenue correcte, propre et professionnelle est exigée en toutes circonstances.",
                "Les vêtements déchirés, trop courts, transparents ou à slogans sont interdits.",
                "Le port des équipements de protection individuelle (EPI) est obligatoire dans les zones concernées.",
                "Toute tenue jugée non conforme pourra entraîner un refus d'accès au poste de travail."
            ]
        },
        {
            "num": "Art. 4",
            "titre": "Respect du silence dans le bâtiment",
            "contenu": [
                "Le bâtiment est un espace partagé avec d'autres services et occupants.",
                "Il est formellement interdit de crier, de parler fort dans les couloirs, salles de réunion et espaces communs.",
                "Les réunions bruyantes doivent se tenir dans les salles prévues à cet effet, portes fermées.",
                "Musique, vidéos ou sonneries audibles depuis les postes de travail sont interdits."
            ]
        },
        {
            "num": "Art. 5",
            "titre": "Hygiène et propreté des lieux",
            "contenu": [
                "Chaque agent est responsable de la propreté de son espace de travail.",
                "Les déchets doivent être déposés dans les poubelles prévues à cet effet.",
                "La cuisine et les espaces communs doivent être laissés propres après utilisation.",
                "Il est interdit de consommer de la nourriture aux postes de travail.",
                "Tout dégât ou désordre constaté doit être signalé sans délai au responsable."
            ]
        },
        {
            "num": "Art. 6",
            "titre": "Sanctions",
            "contenu": [
                "Tout manquement aux dispositions du présent règlement est susceptible de donner lieu à une sanction disciplinaire.",
                "Les sanctions applicables vont de l'avertissement verbal à la mise à pied, selon la gravité et la répétition des faits.",
                "La procédure disciplinaire respectera les droits de la défense de l'agent concerné."
            ]
        }
    ]

    BLEU = RGBColor(0x1A, 0x56, 0x76)

    for regle in regles:
        # Titre article
        h = doc.add_heading("", level=2)
        h.clear()
        r_num = h.add_run(f"{regle['num']} — ")
        r_num.font.color.rgb = BLEU
        r_num.font.size = Pt(13)
        r_num.font.bold = True
        r_titre = h.add_run(regle["titre"])
        r_titre.font.color.rgb = BLEU
        r_titre.font.size = Pt(13)
        r_titre.font.bold = True

        for ligne in regle["contenu"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(ligne).font.size = Pt(11)

        doc.add_paragraph()

    # ── Signatures ──
    doc.add_page_break()
    sig_titre = doc.add_paragraph()
    sig_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_titre.add_run("SIGNATURES").font.bold = True

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Le Responsable Qualité"
    table.cell(0, 1).text = "L'Agent (Lu et approuvé)"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).paragraphs[0].runs[0].bold = True
    table.cell(1, 0).text = "\n\nNom : ___________________"
    table.cell(1, 1).text = "\n\nNom : ___________________"
    table.cell(2, 0).text = "Signature : ______________"
    table.cell(2, 1).text = "Signature : ______________"

    out = EXPORT_DIR / "Reglement_Interieur.docx"
    doc.save(str(out))
    print(f"  ✓ Règlement Intérieur Word : {out.name}")

# ─────────────────────────────────────────────
# 3. FICHIER EXCEL — POINTAGE
# ─────────────────────────────────────────────
def generer_excel_pointage():
    from openpyxl import Workbook
    from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                                  GradientFill)
    from openpyxl.utils import get_column_letter
    from openpyxl.formatting.rule import ColorScaleRule, CellIsRule, FormulaRule
    from openpyxl.styles.differential import DifferentialStyle

    with open(DATA_DIR / "agents.json", encoding="utf-8") as f:
        agents = json.load(f)

    wb = Workbook()

    # ── Couleurs ──
    BLEU_FONCE = "1A5676"
    BLEU_CLAIR = "D0E8F2"
    VERT       = "C6EFCE"
    ROUGE      = "FFC7CE"
    ORANGE     = "FFEB9C"
    GRIS       = "F2F2F2"

    def style_header(cell, bg=BLEU_FONCE, fg="FFFFFF", bold=True, size=11):
        cell.font = Font(bold=bold, color=fg, size=size)
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def border_all(ws, min_row, max_row, min_col, max_col):
        thin = Side(style="thin", color="AAAAAA")
        brd = Border(left=thin, right=thin, top=thin, bottom=thin)
        for row in ws.iter_rows(min_row=min_row, max_row=max_row,
                                 min_col=min_col, max_col=max_col):
            for cell in row:
                cell.border = brd

    # ════════════════════════════════════════
    # FEUILLE 1 — POINTAGE JOURNALIER
    # ════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "Pointage Journalier"
    ws1.sheet_view.showGridLines = False

    # Titre
    ws1.merge_cells("A1:I1")
    ws1["A1"] = "FEUILLE DE POINTAGE JOURNALIER — SERVICE QUALITÉ"
    style_header(ws1["A1"], size=14)
    ws1.row_dimensions[1].height = 30

    # Date
    ws1.merge_cells("A2:I2")
    ws1["A2"] = f"Mois : {datetime.now().strftime('%B %Y').upper()}"
    ws1["A2"].font = Font(bold=True, size=11, color=BLEU_FONCE)
    ws1["A2"].alignment = Alignment(horizontal="center")

    # En-têtes colonnes
    headers = ["ID Agent", "Nom", "Prénom", "Poste",
               "Date", "Heure Arrivée", "Heure Départ", "Statut", "Observations"]
    for col, h in enumerate(headers, 1):
        cell = ws1.cell(row=4, column=col, value=h)
        style_header(cell, bg=BLEU_FONCE)
    ws1.row_dimensions[4].height = 22

    # Données agents (lignes vides pour saisie)
    jours = 26  # ~1 mois de travail
    row = 5
    fills = {
        "Présent": PatternFill("solid", fgColor=VERT),
        "Absent":  PatternFill("solid", fgColor=ROUGE),
        "Retard":  PatternFill("solid", fgColor=ORANGE),
    }
    for agent in agents:
        for j in range(jours):
            ws1.cell(row=row, column=1, value=agent["id"])
            ws1.cell(row=row, column=2, value=agent["nom"])
            ws1.cell(row=row, column=3, value=agent["prenom"])
            ws1.cell(row=row, column=4, value=agent["poste"])
            ws1.cell(row=row, column=5, value=f"Jour {j+1:02d}")
            ws1.cell(row=row, column=6, value="")  # heure arrivée
            ws1.cell(row=row, column=7, value="")  # heure départ
            ws1.cell(row=row, column=8, value="Présent")
            ws1.cell(row=row, column=8).fill = fills["Présent"]
            ws1.cell(row=row, column=9, value="")
            if row % 2 == 0:
                for c in range(1, 10):
                    if ws1.cell(row=row, column=c).fill.fgColor.rgb == "00000000":
                        ws1.cell(row=row, column=c).fill = PatternFill("solid", fgColor=GRIS)
            row += 1

    border_all(ws1, 4, row-1, 1, 9)

    # Largeurs colonnes
    widths = [10, 15, 15, 20, 10, 14, 14, 12, 22]
    for i, w in enumerate(widths, 1):
        ws1.column_dimensions[get_column_letter(i)].width = w

    ws1.freeze_panes = "A5"

    # ════════════════════════════════════════
    # FEUILLE 2 — RÉCAPITULATIF MENSUEL
    # ════════════════════════════════════════
    ws2 = wb.create_sheet("Récapitulatif Mensuel")
    ws2.sheet_view.showGridLines = False

    ws2.merge_cells("A1:H1")
    ws2["A1"] = "RÉCAPITULATIF MENSUEL DES PRÉSENCES"
    style_header(ws2["A1"], size=14)
    ws2.row_dimensions[1].height = 30

    headers2 = ["ID", "Nom", "Prénom", "Poste",
                "Jours Présents", "Jours Absents", "Retards", "Taux Présence (%)"]
    for col, h in enumerate(headers2, 1):
        cell = ws2.cell(row=3, column=col, value=h)
        style_header(cell, bg=BLEU_FONCE)
    ws2.row_dimensions[3].height = 22

    for i, agent in enumerate(agents, 1):
        r = 3 + i
        ws2.cell(row=r, column=1, value=agent["id"])
        ws2.cell(row=r, column=2, value=agent["nom"])
        ws2.cell(row=r, column=3, value=agent["prenom"])
        ws2.cell(row=r, column=4, value=agent["poste"])
        ws2.cell(row=r, column=5, value=0)
        ws2.cell(row=r, column=6, value=0)
        ws2.cell(row=r, column=7, value=0)
        tx = ws2.cell(row=r, column=8,
                      value=f"=IF((E{r}+F{r})=0,0,ROUND(E{r}/(E{r}+F{r})*100,1))")
        tx.number_format = '0.0"%"'
        if r % 2 == 0:
            for c in range(1, 9):
                ws2.cell(row=r, column=c).fill = PatternFill("solid", fgColor=GRIS)

    border_all(ws2, 3, 3 + len(agents), 1, 8)
    for i, w in enumerate([10, 15, 15, 20, 14, 14, 10, 18], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    # ════════════════════════════════════════
    # FEUILLE 3 — REGISTRE AGENTS
    # ════════════════════════════════════════
    ws3 = wb.create_sheet("Registre Agents")
    ws3.sheet_view.showGridLines = False

    ws3.merge_cells("A1:G1")
    ws3["A1"] = "REGISTRE DES AGENTS — SERVICE QUALITÉ"
    style_header(ws3["A1"], size=14)
    ws3.row_dimensions[1].height = 30

    headers3 = ["ID Agent", "Nom", "Prénom", "Poste", "Date d'embauche", "QR Code (fichier)", "Statut"]
    for col, h in enumerate(headers3, 1):
        cell = ws3.cell(row=3, column=col, value=h)
        style_header(cell, bg=BLEU_FONCE)
    ws3.row_dimensions[3].height = 22

    for i, agent in enumerate(agents, 1):
        r = 3 + i
        ws3.cell(row=r, column=1, value=agent["id"])
        ws3.cell(row=r, column=2, value=agent["nom"])
        ws3.cell(row=r, column=3, value=agent["prenom"])
        ws3.cell(row=r, column=4, value=agent["poste"])
        ws3.cell(row=r, column=5, value=agent["date_embauche"])
        ws3.cell(row=r, column=6, value=f"QR_{agent['id']}_{agent['nom']}.png")
        ws3.cell(row=r, column=7, value="Actif")
        ws3.cell(row=r, column=7).fill = PatternFill("solid", fgColor=VERT)
        if r % 2 == 0:
            for c in range(1, 7):
                ws3.cell(row=r, column=c).fill = PatternFill("solid", fgColor=GRIS)

    border_all(ws3, 3, 3 + len(agents), 1, 7)
    for i, w in enumerate([12, 15, 15, 22, 16, 28, 10], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    # ════════════════════════════════════════
    # FEUILLE 4 — STATISTIQUES
    # ════════════════════════════════════════
    ws4 = wb.create_sheet("Statistiques")
    ws4.sheet_view.showGridLines = False

    ws4.merge_cells("A1:D1")
    ws4["A1"] = "TABLEAU DE BORD — STATISTIQUES QUALITÉ"
    style_header(ws4["A1"], size=14)
    ws4.row_dimensions[1].height = 30

    stats = [
        ("", ""),
        ("INDICATEURS DE PRÉSENCE", ""),
        ("Total agents actifs", f"=COUNTA('Registre Agents'!A4:A1000)"),
        ("Jours ouvrés du mois", 26),
        ("", ""),
        ("RÈGLEMENT INTÉRIEUR", ""),
        ("Infractions téléphone (mois)", 0),
        ("Non-conformités vestimentaires", 0),
        ("Incidents bruit signalés", 0),
        ("Non-conformités hygiène", 0),
        ("", ""),
        ("DERNIÈRE MISE À JOUR", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]

    for i, (label, val) in enumerate(stats, 3):
        ws4.cell(row=i, column=1, value=label)
        ws4.cell(row=i, column=2, value=val)
        if label and label == label.upper() and label != "":
            ws4.cell(row=i, column=1).font = Font(bold=True, color=BLEU_FONCE, size=11)
        elif label:
            ws4.cell(row=i, column=1).fill = PatternFill("solid", fgColor=BLEU_CLAIR)
            ws4.cell(row=i, column=2).fill = PatternFill("solid", fgColor=BLEU_CLAIR)

    ws4.column_dimensions["A"].width = 35
    ws4.column_dimensions["B"].width = 25

    out = EXPORT_DIR / "Gestion_Qualite_Pointage.xlsx"
    wb.save(str(out))
    print(f"  ✓ Fichier Excel : {out.name}")

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("\n=== Génération des documents Gestion Qualité ===\n")

    print("1. QR Codes agents...")
    generer_qrcodes()

    print("\n2. Règlement Intérieur (Word)...")
    generer_reglement_word()

    print("\n3. Fichier Excel (pointage + registre + stats)...")
    generer_excel_pointage()

    print("\n✅ Tous les documents ont été générés dans le dossier 'exports/' et 'qrcodes/'")
    print(f"   Répertoire : {BASE_DIR}")
